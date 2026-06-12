"""
BA.IA — Modulo Rendicontazione v1.0
Gestione completa del ciclo post-approvazione bandi.

Funzionalità:
  - Timeline rendicontazione con milestone configurabili
  - Upload giustificativi (fatture, bonifici, contratti)
  - Check automatico documentazione obbligatoria
  - Calcolo importi rendicontati vs ammessi
  - Generazione report rendicontazione PDF/Word per ente
  - Promemoria automatici scadenze SAL
"""

import os, json, secrets, datetime, io, re
from pathlib import Path
import aiosqlite
from fastapi import Request, UploadFile, File
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from security import is_valid_pdf_bytes, sanitize_filename, MAX_UPLOAD_BYTES

DB = lambda: os.environ.get("DB_PATH", "./data/ai-bandi.db")
DOCS_DIR = lambda: Path(DB()).parent / "rendicontazione_docs"


async def init_rendicontazione_db():
    DOCS_DIR().mkdir(parents=True, exist_ok=True)
    async with aiosqlite.connect(DB()) as db:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS rendicontazioni (
                id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                sal_id TEXT,
                bando_id TEXT,
                azienda_id TEXT,
                titolo TEXT NOT NULL,
                importo_approvato REAL DEFAULT 0,
                importo_rendicontato REAL DEFAULT 0,
                data_approvazione TEXT,
                data_scadenza_sal TEXT,
                tipo_sal TEXT DEFAULT 'unico',
                stato TEXT DEFAULT 'aperta',
                ente_erogatore TEXT,
                portale_ente TEXT,
                note TEXT,
                config TEXT,
                created_at TEXT DEFAULT (datetime('now')),
                updated_at TEXT DEFAULT (datetime('now'))
            )""")

        await db.execute("""
            CREATE TABLE IF NOT EXISTS rendicontazione_milestones (
                id TEXT PRIMARY KEY,
                rendicontazione_id TEXT NOT NULL,
                titolo TEXT NOT NULL,
                descrizione TEXT,
                scadenza TEXT,
                ordine INTEGER DEFAULT 0,
                stato TEXT DEFAULT 'pending',
                importo_atteso REAL DEFAULT 0,
                importo_rendicontato REAL DEFAULT 0,
                completata_il TEXT,
                created_at TEXT DEFAULT (datetime('now'))
            )""")

        await db.execute("""
            CREATE TABLE IF NOT EXISTS rendicontazione_documenti (
                id TEXT PRIMARY KEY,
                rendicontazione_id TEXT NOT NULL,
                milestone_id TEXT,
                tipo TEXT NOT NULL,
                titolo TEXT NOT NULL,
                fornitore TEXT,
                numero_documento TEXT,
                data_documento TEXT,
                importo_imponibile REAL DEFAULT 0,
                importo_iva REAL DEFAULT 0,
                importo_totale REAL DEFAULT 0,
                importo_ammissibile REAL DEFAULT 0,
                filename TEXT,
                size_bytes INTEGER DEFAULT 0,
                spesa_categoria TEXT,
                modalita_pagamento TEXT,
                pagato INTEGER DEFAULT 0,
                data_pagamento TEXT,
                note TEXT,
                ai_extracted INTEGER DEFAULT 0,
                ai_confidence TEXT,
                created_at TEXT DEFAULT (datetime('now'))
            )""")

        await db.execute("""
            CREATE TABLE IF NOT EXISTS rendicontazione_checklist (
                id TEXT PRIMARY KEY,
                rendicontazione_id TEXT NOT NULL,
                testo TEXT NOT NULL,
                obbligatorio INTEGER DEFAULT 1,
                completato INTEGER DEFAULT 0,
                ordine INTEGER DEFAULT 0,
                note TEXT,
                created_at TEXT DEFAULT (datetime('now'))
            )""")

        await db.commit()
    print("[RENDICONT] Tabelle rendicontazione pronte")


def _uid(): return secrets.token_hex(10)
def _now(): return datetime.datetime.utcnow().isoformat()


async def _resolve_user(request: Request):
    token = request.headers.get("X-Auth-Token") or request.cookies.get("baia_token")
    if not token: return None
    async with aiosqlite.connect(DB()) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute(
            "SELECT u.* FROM users u JOIN sessions s ON s.user_id=u.id "
            "WHERE s.token=? AND s.expires_at>datetime('now')", (token,)
        ) as c:
            row = await c.fetchone()
    return dict(row) if row else None


def _default_milestones(tipo_sal: str = "unico") -> list:
    """Genera milestone tipiche per tipo di SAL."""
    if tipo_sal == "doppio":
        return [
            {"titolo": "SAL 1 — Avvio attività", "ordine": 1, "importo_atteso": 0.4},
            {"titolo": "SAL 2 — Conclusione progetto", "ordine": 2, "importo_atteso": 0.6},
        ]
    elif tipo_sal == "triplo":
        return [
            {"titolo": "SAL 1 — Avvio (30%)", "ordine": 1, "importo_atteso": 0.3},
            {"titolo": "SAL 2 — Intermedio (40%)", "ordine": 2, "importo_atteso": 0.4},
            {"titolo": "SAL 3 — Finale (30%)", "ordine": 3, "importo_atteso": 0.3},
        ]
    else:
        return [
            {"titolo": "SAL unico — Rendicontazione finale", "ordine": 1, "importo_atteso": 1.0},
        ]


def _default_checklist() -> list:
    """Checklist standard rendicontazione finanza agevolata italiana."""
    return [
        {"testo": "Tutte le fatture sono intestate al beneficiario", "obbligatorio": True, "ordine": 1},
        {"testo": "I bonifici riportano causale con riferimento al bando", "obbligatorio": True, "ordine": 2},
        {"testo": "Le date dei documenti rientrano nel periodo di ammissibilità", "obbligatorio": True, "ordine": 3},
        {"testo": "I beni acquistati sono presenti nella sede dichiarata", "obbligatorio": True, "ordine": 4},
        {"testo": "Foto/documentazione dei beni acquistati allegata", "obbligatorio": False, "ordine": 5},
        {"testo": "Relazione tecnica finale firmata dal legale rappresentante", "obbligatorio": True, "ordine": 6},
        {"testo": "Quietanze di pagamento per ogni fattura", "obbligatorio": True, "ordine": 7},
        {"testo": "Estratti conto bancari del periodo", "obbligatorio": True, "ordine": 8},
        {"testo": "Dichiarazione DURC regolare", "obbligatorio": True, "ordine": 9},
        {"testo": "Dichiarazione antiriciclaggio (se importo >5000€)", "obbligatorio": False, "ordine": 10},
        {"testo": "Verifica regime aiuti de minimis aggiornato", "obbligatorio": True, "ordine": 11},
        {"testo": "Documentazione obblighi di pubblicità (loghi UE/Stato/Regione)", "obbligatorio": False, "ordine": 12},
    ]


def register_rendicontazione_endpoints(app):

    @app.post("/rendicontazione/create")
    async def create_rendicontazione(request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)

        body = await request.json()
        titolo = (body.get("titolo") or "").strip()
        if not titolo: return JSONResponse({"ok": False, "error": "Titolo obbligatorio"}, status_code=400)

        rid = _uid()
        tipo_sal = body.get("tipo_sal", "unico")
        if tipo_sal not in ("unico", "doppio", "triplo"):
            tipo_sal = "unico"
        importo_approvato = float(body.get("importo_approvato") or 0)
        if importo_approvato < 0:
            return JSONResponse({"ok": False, "error": "Importo approvato non può essere negativo"}, status_code=400)

        async with aiosqlite.connect(DB()) as db:
            await db.execute(
                "INSERT INTO rendicontazioni (id,user_id,sal_id,bando_id,azienda_id,titolo,"
                "importo_approvato,data_approvazione,data_scadenza_sal,tipo_sal,ente_erogatore,"
                "portale_ente,note) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (rid, user["id"], body.get("sal_id"), body.get("bando_id"), body.get("azienda_id"),
                 titolo, importo_approvato, body.get("data_approvazione"), body.get("data_scadenza_sal"),
                 tipo_sal, body.get("ente_erogatore", ""), body.get("portale_ente", ""), body.get("note", ""))
            )
            # Crea milestones default
            for m in _default_milestones(tipo_sal):
                mid = _uid()
                await db.execute(
                    "INSERT INTO rendicontazione_milestones (id,rendicontazione_id,titolo,ordine,importo_atteso) "
                    "VALUES (?,?,?,?,?)",
                    (mid, rid, m["titolo"], m["ordine"], importo_approvato * m["importo_atteso"])
                )
            # Crea checklist default
            for c in _default_checklist():
                cid = _uid()
                await db.execute(
                    "INSERT INTO rendicontazione_checklist (id,rendicontazione_id,testo,obbligatorio,ordine) "
                    "VALUES (?,?,?,?,?)",
                    (cid, rid, c["testo"], int(c["obbligatorio"]), c["ordine"])
                )
            await db.commit()

        return {"ok": True, "id": rid}

    @app.get("/rendicontazione/list")
    async def list_rendicontazioni(request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)

        async with aiosqlite.connect(DB()) as db:
            db.row_factory = aiosqlite.Row
            async with db.execute(
                "SELECT * FROM rendicontazioni WHERE user_id=? ORDER BY created_at DESC",
                (user["id"],)
            ) as c:
                rows = [dict(r) for r in await c.fetchall()]

            # Arricchisci con conteggi
            for r in rows:
                async with db.execute(
                    "SELECT COUNT(*) as n, SUM(importo_totale) as t FROM rendicontazione_documenti WHERE rendicontazione_id=?",
                    (r["id"],)
                ) as c:
                    doc_row = await c.fetchone()
                    r["doc_count"] = doc_row[0] or 0
                    r["total_docs_amount"] = doc_row[1] or 0
                async with db.execute(
                    "SELECT COUNT(*) FROM rendicontazione_checklist WHERE rendicontazione_id=? AND completato=1",
                    (r["id"],)
                ) as c:
                    r["checklist_done"] = (await c.fetchone())[0] or 0
                async with db.execute(
                    "SELECT COUNT(*) FROM rendicontazione_checklist WHERE rendicontazione_id=?",
                    (r["id"],)
                ) as c:
                    r["checklist_total"] = (await c.fetchone())[0] or 0

        return {"ok": True, "rendicontazioni": rows}

    @app.get("/rendicontazione/{rid}")
    async def get_rendicontazione(rid: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)

        async with aiosqlite.connect(DB()) as db:
            db.row_factory = aiosqlite.Row
            async with db.execute(
                "SELECT * FROM rendicontazioni WHERE id=? AND user_id=?", (rid, user["id"])
            ) as c:
                r = await c.fetchone()
            if not r: return JSONResponse({"ok": False, "error": "Non trovato"}, status_code=404)

            async with db.execute(
                "SELECT * FROM rendicontazione_milestones WHERE rendicontazione_id=? ORDER BY ordine", (rid,)
            ) as c:
                milestones = [dict(m) for m in await c.fetchall()]

            async with db.execute(
                "SELECT * FROM rendicontazione_documenti WHERE rendicontazione_id=? ORDER BY data_documento DESC, created_at DESC", (rid,)
            ) as c:
                docs = [dict(d) for d in await c.fetchall()]

            async with db.execute(
                "SELECT * FROM rendicontazione_checklist WHERE rendicontazione_id=? ORDER BY ordine", (rid,)
            ) as c:
                checklist = [dict(i) for i in await c.fetchall()]

            # Calcola importo per milestone
            for m in milestones:
                async with db.execute(
                    "SELECT SUM(importo_ammissibile) FROM rendicontazione_documenti WHERE milestone_id=?",
                    (m["id"],)
                ) as c:
                    m["importo_rendicontato_real"] = (await c.fetchone())[0] or 0

        return {"ok": True, "rendicontazione": dict(r),
                "milestones": milestones, "documenti": docs, "checklist": checklist}

    @app.delete("/rendicontazione/{rid}")
    async def delete_rendicontazione(rid: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)
        async with aiosqlite.connect(DB()) as db:
            await db.execute("DELETE FROM rendicontazione_documenti WHERE rendicontazione_id=?", (rid,))
            await db.execute("DELETE FROM rendicontazione_milestones WHERE rendicontazione_id=?", (rid,))
            await db.execute("DELETE FROM rendicontazione_checklist WHERE rendicontazione_id=?", (rid,))
            await db.execute("DELETE FROM rendicontazioni WHERE id=? AND user_id=?", (rid, user["id"]))
            await db.commit()
        return {"ok": True}

    @app.post("/rendicontazione/{rid}/documento")
    async def upload_documento(rid: str, request: Request, file: UploadFile = File(...)):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)

        async with aiosqlite.connect(DB()) as db:
            async with db.execute("SELECT id FROM rendicontazioni WHERE id=? AND user_id=?", (rid, user["id"])) as c:
                if not await c.fetchone():
                    return JSONResponse({"ok": False, "error": "Non trovato"}, status_code=404)

        content = await file.read()
        size = len(content)

        if size > MAX_UPLOAD_BYTES:
            return JSONResponse({"ok": False, "error": f"File troppo grande (max {MAX_UPLOAD_BYTES // 1024 // 1024} MB)"}, status_code=400)
        if not file.filename.lower().endswith(".pdf"):
            return JSONResponse({"ok": False, "error": "Solo file PDF supportati"}, status_code=400)
        if not is_valid_pdf_bytes(content):
            return JSONResponse({"ok": False, "error": "Il file non è un PDF valido"}, status_code=400)

        doc_id    = _uid()
        safe_name = sanitize_filename(file.filename, doc_id)
        doc_dir   = DOCS_DIR() / rid
        doc_dir.mkdir(parents=True, exist_ok=True)
        (doc_dir / safe_name).write_bytes(content)

        # Estrai dati con AI se è un PDF
        ai_data = {}
        ai_confidence = "media"
        if file.filename.lower().endswith(".pdf"):
            try:
                tmp_path = doc_dir / safe_name
                from main import extract_text
                text = extract_text(str(tmp_path))[:6000]
                if len(text) > 200:
                    from app_locale import ai_call_multi
                    prompt = (
                        "Sei un esperto di rendicontazione bandi italiani. Da questa fattura/documento, estrai "
                        "in JSON (zero altro testo): "
                        '{"tipo":"fattura|bonifico|contratto|preventivo|altro",'
                        '"fornitore":"","numero":"","data":"YYYY-MM-DD",'
                        '"imponibile":0,"iva":0,"totale":0,'
                        '"descrizione":"breve descrizione",'
                        '"modalita_pagamento":"bonifico|carta|contanti|altro",'
                        '"confidence":"alta|media|bassa"}\n\n' + text
                    )
                    result, _ = await ai_call_multi(prompt, json_mode=True, timeout=60)
                    clean = result.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
                    ai_data = json.loads(clean)
                    ai_confidence = ai_data.get("confidence", "media")
            except Exception as e:
                print(f"[RENDICONT] AI extract error: {e}")

        body = dict(request.query_params)
        milestone_id = body.get("milestone_id", "")
        tipo = body.get("tipo") or ai_data.get("tipo") or "fattura"
        titolo = body.get("titolo") or ai_data.get("descrizione") or file.filename

        async with aiosqlite.connect(DB()) as db:
            await db.execute(
                "INSERT INTO rendicontazione_documenti (id,rendicontazione_id,milestone_id,tipo,titolo,"
                "fornitore,numero_documento,data_documento,importo_imponibile,importo_iva,importo_totale,"
                "importo_ammissibile,filename,size_bytes,spesa_categoria,modalita_pagamento,note,"
                "ai_extracted,ai_confidence) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                (doc_id, rid, milestone_id or None, tipo, titolo,
                 ai_data.get("fornitore", ""), ai_data.get("numero", ""), ai_data.get("data"),
                 float(ai_data.get("imponibile") or 0), float(ai_data.get("iva") or 0),
                 float(ai_data.get("totale") or 0), float(ai_data.get("totale") or 0),
                 file.filename, size, body.get("categoria", ""),
                 ai_data.get("modalita_pagamento", ""), body.get("note", ""),
                 1 if ai_data else 0, ai_confidence)
            )
            # Aggiorna importo_rendicontato totale
            await db.execute(
                "UPDATE rendicontazioni SET importo_rendicontato="
                "(SELECT COALESCE(SUM(importo_ammissibile),0) FROM rendicontazione_documenti WHERE rendicontazione_id=?),"
                "updated_at=? WHERE id=?",
                (rid, _now(), rid)
            )
            await db.commit()

        return {"ok": True, "id": doc_id, "ai_extracted": bool(ai_data), "data": ai_data}

    @app.put("/rendicontazione/documento/{doc_id}")
    async def update_documento(doc_id: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)

        body = await request.json()
        _ALLOWED_FIELDS = {
            "tipo", "titolo", "fornitore", "numero_documento", "data_documento",
            "importo_imponibile", "importo_iva", "importo_totale", "importo_ammissibile",
            "spesa_categoria", "modalita_pagamento", "pagato", "data_pagamento", "note",
            "milestone_id"
        }
        updates = []
        params = []
        for f in _ALLOWED_FIELDS:
            if f in body:
                updates.append(f"{f}=?")
                params.append(body[f])

        if not updates: return {"ok": True, "message": "Nessuna modifica"}

        async with aiosqlite.connect(DB()) as db:
            # Verifica che il documento appartenga a una rendicontazione dell'utente
            async with db.execute(
                "SELECT rd.rendicontazione_id FROM rendicontazione_documenti rd "
                "JOIN rendicontazioni r ON r.id=rd.rendicontazione_id "
                "WHERE rd.id=? AND r.user_id=?",
                (doc_id, user["id"])
            ) as c:
                row = await c.fetchone()
            if not row:
                return JSONResponse({"ok": False, "error": "Non trovato o non autorizzato"}, status_code=404)

            rid = row[0]
            params.append(doc_id)
            await db.execute(f"UPDATE rendicontazione_documenti SET {','.join(updates)} WHERE id=?", params)
            await db.execute(
                "UPDATE rendicontazioni SET importo_rendicontato="
                "(SELECT COALESCE(SUM(importo_ammissibile),0) FROM rendicontazione_documenti WHERE rendicontazione_id=?),"
                "updated_at=? WHERE id=?",
                (rid, _now(), rid)
            )
            await db.commit()
        return {"ok": True}

    @app.delete("/rendicontazione/documento/{doc_id}")
    async def delete_documento(doc_id: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)
        async with aiosqlite.connect(DB()) as db:
            # Verifica ownership tramite JOIN
            async with db.execute(
                "SELECT rd.rendicontazione_id FROM rendicontazione_documenti rd "
                "JOIN rendicontazioni r ON r.id=rd.rendicontazione_id "
                "WHERE rd.id=? AND r.user_id=?",
                (doc_id, user["id"])
            ) as c:
                row = await c.fetchone()
            if not row:
                return JSONResponse({"ok": False, "error": "Non trovato o non autorizzato"}, status_code=404)
            rid = row[0]
            await db.execute("DELETE FROM rendicontazione_documenti WHERE id=?", (doc_id,))
            await db.execute(
                "UPDATE rendicontazioni SET importo_rendicontato="
                "(SELECT COALESCE(SUM(importo_ammissibile),0) FROM rendicontazione_documenti WHERE rendicontazione_id=?) "
                "WHERE id=?", (rid, rid)
            )
            await db.commit()
        return {"ok": True}

    @app.put("/rendicontazione/milestone/{mid}")
    async def update_milestone(mid: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)
        body = await request.json()
        async with aiosqlite.connect(DB()) as db:
            # Verifica ownership
            async with db.execute(
                "SELECT m.id FROM rendicontazione_milestones m "
                "JOIN rendicontazioni r ON r.id=m.rendicontazione_id "
                "WHERE m.id=? AND r.user_id=?", (mid, user["id"])
            ) as c:
                if not await c.fetchone():
                    return JSONResponse({"ok": False, "error": "Non trovato o non autorizzato"}, status_code=404)
            await db.execute(
                "UPDATE rendicontazione_milestones SET titolo=COALESCE(?,titolo),"
                "descrizione=COALESCE(?,descrizione),scadenza=COALESCE(?,scadenza),"
                "stato=COALESCE(?,stato),importo_atteso=COALESCE(?,importo_atteso),"
                "completata_il=CASE WHEN ?='completata' THEN ? ELSE completata_il END "
                "WHERE id=?",
                (body.get("titolo"), body.get("descrizione"), body.get("scadenza"),
                 body.get("stato"), body.get("importo_atteso"),
                 body.get("stato"), _now(), mid)
            )
            await db.commit()
        return {"ok": True}

    @app.put("/rendicontazione/checklist/{cid}")
    async def update_checklist_item(cid: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)
        body = await request.json()
        async with aiosqlite.connect(DB()) as db:
            # Verifica ownership
            async with db.execute(
                "SELECT c.id FROM rendicontazione_checklist c "
                "JOIN rendicontazioni r ON r.id=c.rendicontazione_id "
                "WHERE c.id=? AND r.user_id=?", (cid, user["id"])
            ) as c2:
                if not await c2.fetchone():
                    return JSONResponse({"ok": False, "error": "Non trovato o non autorizzato"}, status_code=404)
            await db.execute(
                "UPDATE rendicontazione_checklist SET completato=COALESCE(?,completato),"
                "testo=COALESCE(?,testo),note=COALESCE(?,note) WHERE id=?",
                (body.get("completato"), body.get("testo"), body.get("note"), cid)
            )
            await db.commit()
        return {"ok": True}

    @app.post("/rendicontazione/{rid}/checklist")
    async def add_checklist_item(rid: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)
        body = await request.json()
        async with aiosqlite.connect(DB()) as db:
            # Verifica ownership della rendicontazione
            async with db.execute(
                "SELECT id FROM rendicontazioni WHERE id=? AND user_id=?", (rid, user["id"])
            ) as c:
                if not await c.fetchone():
                    return JSONResponse({"ok": False, "error": "Non trovato o non autorizzato"}, status_code=404)
            cid = _uid()
            await db.execute(
                "INSERT INTO rendicontazione_checklist (id,rendicontazione_id,testo,obbligatorio,ordine) "
                "VALUES (?,?,?,?,?)",
                (cid, rid, body.get("testo", ""), int(body.get("obbligatorio", 0)), int(body.get("ordine", 99)))
            )
            await db.commit()
        return {"ok": True, "id": cid}

    @app.delete("/rendicontazione/checklist/{cid}")
    async def delete_checklist_item(cid: str, request: Request):
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)
        async with aiosqlite.connect(DB()) as db:
            # Verifica ownership
            async with db.execute(
                "SELECT c.id FROM rendicontazione_checklist c "
                "JOIN rendicontazioni r ON r.id=c.rendicontazione_id "
                "WHERE c.id=? AND r.user_id=?", (cid, user["id"])
            ) as c2:
                if not await c2.fetchone():
                    return JSONResponse({"ok": False, "error": "Non trovato o non autorizzato"}, status_code=404)
            await db.execute("DELETE FROM rendicontazione_checklist WHERE id=?", (cid,))
            await db.commit()
        return {"ok": True}

    @app.get("/rendicontazione/{rid}/export/word")
    async def export_rendicontazione_word(rid: str, request: Request):
        """Genera report rendicontazione professionale .docx."""
        user = await _resolve_user(request)
        if not user: return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)

        async with aiosqlite.connect(DB()) as db:
            db.row_factory = aiosqlite.Row
            async with db.execute("SELECT * FROM rendicontazioni WHERE id=? AND user_id=?", (rid, user["id"])) as c:
                r = await c.fetchone()
            if not r: return JSONResponse({"ok": False, "error": "Non trovato"}, status_code=404)
            r = dict(r)
            async with db.execute("SELECT * FROM rendicontazione_milestones WHERE rendicontazione_id=? ORDER BY ordine", (rid,)) as c:
                milestones = [dict(m) for m in await c.fetchall()]
            async with db.execute("SELECT * FROM rendicontazione_documenti WHERE rendicontazione_id=? ORDER BY data_documento", (rid,)) as c:
                docs = [dict(d) for d in await c.fetchall()]
            async with db.execute("SELECT * FROM rendicontazione_checklist WHERE rendicontazione_id=? ORDER BY ordine", (rid,)) as c:
                checklist = [dict(i) for i in await c.fetchall()]

        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def _bg(cell, hex_color):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        doc = Document()
        for s in doc.sections:
            s.top_margin = s.bottom_margin = Cm(2.0)
            s.left_margin = s.right_margin = Cm(2.2)

        # Header
        p = doc.add_paragraph()
        run = p.add_run("BA.IA")
        run.font.name = 'Calibri'; run.font.size = Pt(20); run.font.bold = True
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x14)
        sub = doc.add_paragraph()
        sub.add_run("RAPPORTO DI RENDICONTAZIONE").font.color.rgb = RGBColor(0xC8, 0xA8, 0x4B)
        doc.add_paragraph()

        # Titolo
        title_p = doc.add_paragraph()
        tr = title_p.add_run(r["titolo"])
        tr.font.name = 'Calibri'; tr.font.size = Pt(15); tr.font.bold = True

        # Tabella info
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_data = [
            ("Ente erogatore", r.get("ente_erogatore") or "—"),
            ("Tipo SAL", r.get("tipo_sal", "unico").upper()),
            ("Data approvazione", r.get("data_approvazione") or "—"),
            ("Scadenza SAL", r.get("data_scadenza_sal") or "—"),
            ("Importo approvato", f"€ {r['importo_approvato']:,.2f}".replace(",", ".")),
            ("Importo rendicontato", f"€ {r['importo_rendicontato']:,.2f}".replace(",", ".")),
            ("Stato", r.get("stato", "aperta").upper()),
        ]
        for label, value in info_data:
            row = info_table.add_row()
            row.cells[0].paragraphs[0].add_run(label).font.bold = True
            row.cells[1].paragraphs[0].add_run(str(value))
            _bg(row.cells[0], 'F4F4F8')

        doc.add_paragraph()

        # Milestone
        if milestones:
            h = doc.add_paragraph().add_run("MILESTONE")
            h.font.size = Pt(11); h.font.bold = True; h.font.color.rgb = RGBColor(0xC8, 0xA8, 0x4B)
            ms_table = doc.add_table(rows=1, cols=4)
            ms_table.style = 'Light Grid Accent 1'
            hdr = ms_table.rows[0].cells
            for i, t in enumerate(["Milestone", "Scadenza", "Atteso", "Rendicontato"]):
                p = hdr[i].paragraphs[0]; p.add_run(t).font.bold = True
                _bg(hdr[i], 'EBE2C0')
            for m in milestones:
                row = ms_table.add_row().cells
                row[0].paragraphs[0].add_run(m["titolo"])
                row[1].paragraphs[0].add_run(m.get("scadenza") or "—")
                row[2].paragraphs[0].add_run(f"€ {m['importo_atteso']:,.2f}".replace(",", "."))
                # Calcola rendicontato per milestone
                ms_docs = [d for d in docs if d.get("milestone_id") == m["id"]]
                rend = sum(d.get("importo_ammissibile", 0) for d in ms_docs)
                row[3].paragraphs[0].add_run(f"€ {rend:,.2f}".replace(",", "."))
            doc.add_paragraph()

        # Documenti rendicontati
        if docs:
            h = doc.add_paragraph().add_run("DOCUMENTI RENDICONTATI")
            h.font.size = Pt(11); h.font.bold = True; h.font.color.rgb = RGBColor(0xC8, 0xA8, 0x4B)
            doc_table = doc.add_table(rows=1, cols=6)
            doc_table.style = 'Light Grid Accent 1'
            hdr = doc_table.rows[0].cells
            for i, t in enumerate(["Tipo", "Fornitore", "N. doc", "Data", "Importo", "Ammissibile"]):
                hdr[i].paragraphs[0].add_run(t).font.bold = True
                _bg(hdr[i], 'EBE2C0')
            for d in docs:
                row = doc_table.add_row().cells
                row[0].paragraphs[0].add_run(d.get("tipo", ""))
                row[1].paragraphs[0].add_run((d.get("fornitore") or "")[:30])
                row[2].paragraphs[0].add_run(d.get("numero_documento") or "—")
                row[3].paragraphs[0].add_run(d.get("data_documento") or "—")
                row[4].paragraphs[0].add_run(f"€ {d.get('importo_totale',0):,.2f}".replace(",", "."))
                row[5].paragraphs[0].add_run(f"€ {d.get('importo_ammissibile',0):,.2f}".replace(",", "."))

            # Totale
            tot_row = doc_table.add_row().cells
            tot_row[0].merge(tot_row[3]).paragraphs[0].add_run("TOTALE").bold = True
            tot_imp = sum(d.get("importo_totale", 0) for d in docs)
            tot_amm = sum(d.get("importo_ammissibile", 0) for d in docs)
            tot_row[4].paragraphs[0].add_run(f"€ {tot_imp:,.2f}".replace(",", ".")).bold = True
            tot_row[5].paragraphs[0].add_run(f"€ {tot_amm:,.2f}".replace(",", ".")).bold = True
            for c in tot_row:
                _bg(c, 'F4F4F8')

            doc.add_paragraph()

        # Checklist conformità
        if checklist:
            h = doc.add_paragraph().add_run("CHECKLIST CONFORMITÀ")
            h.font.size = Pt(11); h.font.bold = True; h.font.color.rgb = RGBColor(0xC8, 0xA8, 0x4B)
            for item in checklist:
                p = doc.add_paragraph(style='List Bullet')
                symbol = "✓ " if item["completato"] else "○ "
                run = p.add_run(symbol + item["testo"])
                if item["completato"]:
                    run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
                elif item["obbligatorio"]:
                    run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                run.font.size = Pt(9)

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        f = footer.add_run(f"Rapporto generato da BA.IA il {datetime.datetime.now().strftime('%d/%m/%Y alle %H:%M')} — Da verificare con la documentazione ufficiale dell'ente erogatore.")
        f.font.size = Pt(7); f.font.color.rgb = RGBColor(0xAA, 0xAA, 0xCC); f.italic = True

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        safe_name = re.sub(r"[^a-zA-Z0-9_\-]", "_", r["titolo"][:50])
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="Rendicontazione_{safe_name}.docx"'})

    # ── A3.1 INVOICE PROCESSOR OCR ─────────────────────────
    @app.post("/rendicontazione/ocr-fattura")
    async def ocr_fattura(request: Request, file: UploadFile = File(...)):
        """Estrai dati strutturati da fattura PDF tramite text-layer + AI."""
        user = await _resolve_user(request)
        if not user:
            return JSONResponse({"ok": False, "error": "Non autenticato"}, status_code=401)

        content = await file.read()
        if len(content) > MAX_UPLOAD_BYTES:
            return JSONResponse({"ok": False, "error": "File troppo grande (max 10MB)"}, status_code=413)
        if not is_valid_pdf_bytes(content):
            return JSONResponse({"ok": False, "error": "File non è un PDF valido"}, status_code=400)

        # Estrai testo dal PDF
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            return JSONResponse({"ok": False, "error": f"Errore lettura PDF: {e}"}, status_code=422)

        if len(text.strip()) < 20:
            return JSONResponse({"ok": False, "error": "PDF senza testo estraibile (PDF immagine — OCR non disponibile in questa versione)"}, status_code=422)

        # Regex patterns per fatture italiane
        def _find(patterns, text):
            for p in patterns:
                m = re.search(p, text, re.IGNORECASE | re.MULTILINE)
                if m:
                    return m.group(1).strip()
            return None

        numero = _find([
            r"(?:fattura|ft\.?|n\.?)\s*(?:nr\.?|n\.?|numero)?\s*[:\s]?\s*([A-Z0-9/\-_]+)",
            r"(?:documento|doc\.?)\s*n\.?\s*[:\s]?\s*([A-Z0-9/\-_]+)",
        ], text)

        data_fattura = _find([
            r"(?:data\s*(?:fattura|emissione|documento)?)\s*[:\s]?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
            r"(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{4})",
        ], text)

        importo_totale = _find([
            r"(?:totale\s*(?:fattura|documento|da\s*pagare|imponibile\s*\+\s*iva)?)\s*[:\€\s]?\s*([\d\.,]+)",
            r"(?:totale)\s*[€\s]+\s*([\d\.,]+)",
            r"(?:importo\s*totale)\s*[:\€\s]+\s*([\d\.,]+)",
        ], text)

        imponibile = _find([
            r"(?:imponibile|base\s*imponibile)\s*[:\€\s]+\s*([\d\.,]+)",
        ], text)

        iva_percent = _find([
            r"(?:iva|i\.v\.a\.)\s*(\d{1,2})\s*%",
            r"(\d{1,2})\s*%\s*iva",
        ], text)

        iva_importo = _find([
            r"(?:iva|imposta)\s*[:\€\s]+\s*([\d\.,]+)",
        ], text)

        piva_fornitore = _find([
            r"(?:p\.?\s*iva|partita\s*iva)\s*[:\s]+\s*([IT]?\d{11})",
            r"\bIT(\d{11})\b",
        ], text)

        cf_fornitore = _find([
            r"(?:cod\.?\s*fisc\.?|codice\s*fiscale)\s*[:\s]+\s*([A-Z0-9]{11,16})",
        ], text)

        fornitore = _find([
            r"^([A-Z][A-Z\s\.\,]{5,60}(?:SRL|SPA|SNC|SAS|SRLS|SOC\.|LTD|S\.R\.L\.|S\.P\.A\.))",
            r"(?:fornitore|emittente|cedente)\s*[:\s]+\s*(.{5,60})",
        ], text)

        descrizione = _find([
            r"(?:descrizione|oggetto|prestazione)\s*[:\s]+\s*(.{10,120})",
        ], text)

        # Normalizza importi: 1.234,56 → 1234.56
        def _norm_amount(s):
            if not s: return None
            s = s.replace(".", "").replace(",", ".")
            try: return round(float(s), 2)
            except: return None

        result = {
            "numero_fattura": numero,
            "data_fattura": data_fattura,
            "fornitore": fornitore,
            "piva_fornitore": piva_fornitore,
            "cf_fornitore": cf_fornitore,
            "imponibile": _norm_amount(imponibile),
            "iva_percent": iva_percent,
            "iva_importo": _norm_amount(iva_importo),
            "importo_totale": _norm_amount(importo_totale),
            "descrizione": descrizione,
            "testo_estratto_chars": len(text),
            "metodo": "regex",
            "confidence": "alta" if (numero and importo_totale) else "media" if importo_totale else "bassa",
        }

        # Se ANTHROPIC disponibile, arricchisci con AI
        import os
        api_key = os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("AI_API_KEY", "")
        if api_key and len(text) < 8000:
            try:
                import httpx
                ai_prompt = (
                    "Estrai i seguenti campi da questa fattura italiana e rispondimi SOLO con JSON valido:\n"
                    "numero_fattura, data_fattura (YYYY-MM-DD), fornitore (ragione sociale), "
                    "piva_fornitore, imponibile (float), iva_percent (int), iva_importo (float), "
                    "importo_totale (float), descrizione (breve descrizione servizi/beni).\n"
                    "Se un campo non è presente metti null.\n\n"
                    f"FATTURA:\n{text[:4000]}"
                )
                async with httpx.AsyncClient(timeout=30) as client:
                    r = await client.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={"x-api-key": api_key, "anthropic-version": "2023-06-01", "Content-Type": "application/json"},
                        json={"model": "claude-haiku-4-5-20251001", "max_tokens": 512,
                              "system": "Rispondi SEMPRE con JSON valido e nient'altro.",
                              "messages": [{"role": "user", "content": ai_prompt}]}
                    )
                if r.status_code == 200:
                    ai_text = r.json()["content"][0]["text"]
                    ai_json = json.loads(ai_text)
                    for k, v in ai_json.items():
                        if v is not None and result.get(k) is None:
                            result[k] = v
                    result["metodo"] = "regex+ai"
                    result["confidence"] = "alta"
            except Exception:
                pass  # fallback to regex only

        return {"ok": True, "fattura": result}
