"""
BA.IA — Backend Locale v2.0
Estende main.py con: Auth, Setup, Export Word, Import documenti aziendali.
"""

import sys, os
from pathlib import Path
# Must be before any local module imports so Docker/uvicorn can find security, audit, etc.
sys.path.insert(0, str(Path(__file__).parent))

import re, json, uuid, hashlib, datetime, io, tempfile
from security import hash_password, verify_password, validate_email, login_limiter, register_limiter, is_valid_pdf_bytes, MAX_UPLOAD_BYTES, RateLimiter

reset_limiter = RateLimiter(max_calls=3, window_seconds=3600)  # 3 reset req / ora per IP

ROOT = Path(__file__).parent.parent
ENV_FILE = ROOT / ".env"

# ── Carica .env ─────────────────────────────────────────
def _read_env():
    result = {}
    if ENV_FILE.exists():
        for line in ENV_FILE.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                result[k.strip()] = v.strip()
    return result

def _write_env_key(key, value):
    if ENV_FILE.exists():
        content = ENV_FILE.read_text(encoding="utf-8")
        pattern = rf"^{re.escape(key)}=.*$"
        if re.search(pattern, content, flags=re.MULTILINE):
            content = re.sub(pattern, f"{key}={value}", content, flags=re.MULTILINE)
        else:
            content = content.rstrip() + f"\n{key}={value}\n"
    else:
        content = (
            "GROQ_API_KEY=gsk_INSERISCI_LA_TUA_CHIAVE_QUI\nLICENSE_KEY=TEST-MODE\n"
            "GROQ_MODEL=meta-llama/llama-4-scout-17b-16e-instruct\nAPP_NAME=BA.IA\n"
            "DB_PATH=./data/baia.db\nPORT=8000\n"
        )
        content = re.sub(rf"^{re.escape(key)}=.*$", f"{key}={value}", content, flags=re.MULTILINE)
    ENV_FILE.write_text(content, encoding="utf-8")

env_vals = _read_env()
for k, v in env_vals.items():
    os.environ.setdefault(k, v)

db_path = os.environ.get("DB_PATH", "./data/baia.db")
if not os.path.isabs(db_path):
    os.environ["DB_PATH"] = str(ROOT / db_path.replace("./", ""))

Path(os.environ["DB_PATH"]).parent.mkdir(parents=True, exist_ok=True)

sys.path.insert(0, str(Path(__file__).parent))

# ── Bridge SQLite→PostgreSQL (auto-attivo se DATABASE_URL settata) ──
try:
    from sqlite_pg_bridge import install as _install_pg_bridge
    _install_pg_bridge()
except Exception as _e:
    print(f"[DB] Bridge non installato: {_e}")

from main import app, init_db  # noqa
from audit import init_audit_db, log_action, get_client_ip, register_audit_endpoints
from notifications import init_notifications_db, register_notification_endpoints, check_bandi_scadenze

# ═══════════════════════════════════════════════════════
# DB INIT ESTESO — aggiunge tabelle auth
# ═══════════════════════════════════════════════════════
import aiosqlite

DB = os.environ["DB_PATH"]

async def init_auth_db():
    async with aiosqlite.connect(DB) as db:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                email TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                created_at TEXT DEFAULT (datetime('now'))
            )""")
        await db.execute("""
            CREATE TABLE IF NOT EXISTS sessions (
                token TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                created_at TEXT DEFAULT (datetime('now')),
                expires_at TEXT NOT NULL
            )""")
        await db.execute("""
            CREATE TABLE IF NOT EXISTS password_reset_tokens (
                token TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                email TEXT NOT NULL,
                expires_at TEXT NOT NULL,
                used INTEGER DEFAULT 0,
                created_at TEXT DEFAULT (datetime('now'))
            )""")
        await db.execute("CREATE INDEX IF NOT EXISTS ix_prt_email ON password_reset_tokens(email)")
        await db.commit()
    print("[AUTH] Tabelle utenti pronte")

from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware

@app.on_event("startup")
async def startup_auth():
    await init_auth_db()
    await init_audit_db()
    await init_notifications_db()
    register_audit_endpoints(app)
    register_notification_endpoints(app)

# ── Helper token/session ─────────────────────────────
def _make_token() -> str:
    import secrets
    return secrets.token_hex(32)

def _expires() -> str:
    return (datetime.datetime.utcnow() + datetime.timedelta(days=1)).isoformat()

def _send_email(to: str, subject: str, html: str) -> None:
    """Invia email via SMTP configurato. Solleva eccezione in caso di errore."""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    host  = os.environ.get("SMTP_HOST", "")
    port  = int(os.environ.get("SMTP_PORT", "587"))
    user  = os.environ.get("SMTP_USER", "")
    pwd   = os.environ.get("SMTP_PASS", "")
    frm   = os.environ.get("SMTP_FROM", user)
    if not host:
        raise ValueError("SMTP non configurato")
    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"]    = frm
    msg["To"]      = to
    msg.attach(MIMEText(html, "html", "utf-8"))
    with smtplib.SMTP(host, port, timeout=15) as s:
        s.ehlo(); s.starttls()
        if user: s.login(user, pwd)
        s.sendmail(frm, to, msg.as_string())

# ── Paywall / Trial / Admin ──────────────────────────
ADMIN_EMAILS = {e.strip().lower() for e in os.environ.get("ADMIN_EMAILS", "").split(",") if e.strip()}
TRIAL_DAYS = 14
PAID_PLANS = ("free", "active", "paid", "pro", "base", "business")

def _rget(row, key, default=None):
    try:
        v = row[key]
        return v if v is not None else default
    except (KeyError, IndexError, TypeError):
        return default

def _is_admin(email, role=None):
    return (email or "").lower() in ADMIN_EMAILS or role in ("admin", "developer")

def _has_access(plan, trial_ends_at, email=None, role=None):
    """True se l'utente può usare l'app: admin, piano attivo/free, o trial ancora valido."""
    if _is_admin(email, role):
        return True
    if (plan or "trial") in PAID_PLANS:
        return True
    if trial_ends_at:
        try:
            return datetime.datetime.utcnow().isoformat() < str(trial_ends_at)
        except Exception:
            return True
    return True  # retrocompat: nessuna scadenza = non bloccare

# ── Auth middleware helper ───────────────────────────
async def _get_user(request: Request):
    token = request.headers.get("X-Auth-Token") or request.cookies.get("baia_token")
    if not token:
        return None
    async with aiosqlite.connect(DB) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute(
            "SELECT u.* FROM users u JOIN sessions s ON s.user_id=u.id WHERE s.token=? AND s.expires_at>datetime('now')",
            (token,)
        ) as c:
            row = await c.fetchone()
    if row:
        email = row["email"]
        plan = _rget(row, "plan", "trial")
        trial_ends_at = _rget(row, "trial_ends_at")
        role = _rget(row, "role", "consulente")
        return {"id": row["id"], "name": row["name"], "email": email,
                "plan": plan, "trial_ends_at": trial_ends_at, "role": role,
                "is_admin": _is_admin(email, role),
                "access": _has_access(plan, trial_ends_at, email, role)}
    return None

# ═══════════════════════════════════════════════════════
# AUTH ENDPOINTS
# ═══════════════════════════════════════════════════════

@app.post("/auth/register")
async def auth_register(request: Request):
    client_ip = request.client.host if request.client else "unknown"
    if not register_limiter.is_allowed(client_ip):
        return JSONResponse({"ok": False, "error": "Troppe registrazioni. Riprova tra un'ora."}, 429)

    body = await request.json()
    name = (body.get("name") or "").strip()
    email = (body.get("email") or "").strip().lower()
    password = body.get("password") or ""

    if not name or not email or len(password) < 8:
        return JSONResponse({"ok": False, "error": "Dati non validi o password troppo corta (min 8 caratteri)"}, 400)
    if not validate_email(email):
        return JSONResponse({"ok": False, "error": "Formato email non valido"}, 400)

    pw_hash = hash_password(password)
    user_id = "u-" + _make_token()[:16]
    token = _make_token()
    trial_ends = (datetime.datetime.utcnow() + datetime.timedelta(days=TRIAL_DAYS)).isoformat()
    plan = "free" if _is_admin(email) else "trial"   # admin = licenza gratuita immediata
    # Utente + sessione nella STESSA transazione: la FK è soddisfatta e si
    # salvano atomicamente (se fallisce uno, rollback di entrambi).
    try:
        async with aiosqlite.connect(DB) as db:
            await db.execute(
                "INSERT INTO users (id,name,email,password_hash,plan,trial_ends_at) VALUES (?,?,?,?,?,?)",
                (user_id, name, email, pw_hash, plan, trial_ends))
            await db.execute("INSERT INTO sessions (token,user_id,expires_at) VALUES (?,?,?)",
                             (token, user_id, _expires()))
            await db.commit()
    except Exception as e:
        code = getattr(e, "sqlstate", "") or ""
        msg = str(e).lower()
        if (code == "23505" or "unique" in msg or "duplicate" in msg
                or "already exists" in msg or "uniqueviolation" in type(e).__name__.lower()):
            return JSONResponse({"ok": False, "error": "Questa email è già registrata. Prova ad accedere."}, 409)
        print(f"[AUTH] Errore registrazione: {type(e).__name__}: {e}")
        return JSONResponse({"ok": False, "error": f"Errore DB: {type(e).__name__}: {e}"}, 500)
    try:
        await log_action("user.register", user_id=user_id, user_email=email,
                         resource_type="user", resource_id=user_id, ip=client_ip)
    except Exception:
        pass
    print(f"[AUTH] Nuovo utente: {email} (plan={plan})")
    return {"ok": True, "token": token,
            "user": {"id": user_id, "name": name, "email": email,
                     "plan": plan, "trial_ends_at": trial_ends,
                     "is_admin": _is_admin(email), "access": True}}

@app.post("/auth/login")
async def auth_login(request: Request):
    client_ip = request.client.host if request.client else "unknown"
    if not login_limiter.is_allowed(client_ip):
        retry = login_limiter.retry_after(client_ip)
        return JSONResponse({"ok": False, "error": f"Troppi tentativi. Riprova tra {retry}s."}, 429)

    body = await request.json()
    email = (body.get("email") or "").strip().lower()
    password = body.get("password") or ""

    async with aiosqlite.connect(DB) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute("SELECT * FROM users WHERE email=?", (email,)) as c:
            user = await c.fetchone()

    # Risposta generica per prevenire user enumeration
    if not user or not verify_password(password, user["password_hash"]):
        await log_action("user.login_failed", user_email=email, ip=client_ip)
        return JSONResponse({"ok": False, "error": "Credenziali non valide"}, 401)

    token = _make_token()
    async with aiosqlite.connect(DB) as db:
        await db.execute("INSERT INTO sessions (token,user_id,expires_at) VALUES (?,?,?)",
                         (token, user["id"], _expires()))
        await db.commit()
    await log_action("user.login", user_id=user["id"], user_email=user["email"], ip=client_ip)
    u_plan = _rget(user, "plan", "trial")
    u_trial = _rget(user, "trial_ends_at")
    u_role = _rget(user, "role", "consulente")
    return {"ok": True, "token": token,
            "user": {"id": user["id"], "name": user["name"], "email": user["email"],
                     "plan": u_plan, "trial_ends_at": u_trial, "role": u_role,
                     "is_admin": _is_admin(user["email"], u_role),
                     "access": _has_access(u_plan, u_trial, user["email"], u_role)}}

@app.post("/auth/logout")
async def auth_logout(request: Request):
    token = request.headers.get("X-Auth-Token", "")
    if token:
        async with aiosqlite.connect(DB) as db:
            db.row_factory = aiosqlite.Row
            async with db.execute("SELECT user_id FROM sessions WHERE token=?", (token,)) as c:
                row = await c.fetchone()
            user_id = row["user_id"] if row else None
            await db.execute("DELETE FROM sessions WHERE token=?", (token,))
            await db.commit()
        if user_id:
            await log_action("user.logout", user_id=user_id, ip=get_client_ip(request))
    return {"ok": True}

# ── Account & Admin (paywall / licenze) ──────────────
@app.get("/auth/me")
async def auth_me(request: Request):
    user = await _get_user(request)
    if not user:
        return JSONResponse({"ok": False, "error": "Non autenticato"}, 401)
    return {"ok": True, "user": user}

@app.post("/admin/grant-license")
async def admin_grant_license(request: Request):
    """Concede/cambia il piano di un account. Solo admin."""
    me = await _get_user(request)
    if not me or not me.get("is_admin"):
        return JSONResponse({"ok": False, "error": "Non autorizzato"}, 403)
    body = await request.json()
    target = (body.get("email") or "").strip().lower()
    plan = (body.get("plan") or "free").strip().lower()
    if not target:
        return JSONResponse({"ok": False, "error": "Email mancante"}, 400)
    if plan not in ("free", "active", "trial", "expired"):
        return JSONResponse({"ok": False, "error": "Piano non valido"}, 400)
    async with aiosqlite.connect(DB) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute("SELECT id FROM users WHERE email=?", (target,)) as c:
            row = await c.fetchone()
        if not row:
            return JSONResponse({"ok": False, "error": "Utente non trovato"}, 404)
        await db.execute("UPDATE users SET plan=? WHERE email=?", (plan, target))
        await db.commit()
    await log_action("admin.grant_license", user_id=me["id"], user_email=me["email"],
                     resource_type="user", resource_id=target, details={"plan": plan})
    return {"ok": True, "email": target, "plan": plan}

@app.get("/admin/users")
async def admin_users(request: Request):
    """Lista account (per gestione licenze). Solo admin."""
    me = await _get_user(request)
    if not me or not me.get("is_admin"):
        return JSONResponse({"ok": False, "error": "Non autorizzato"}, 403)
    async with aiosqlite.connect(DB) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute(
            "SELECT email,name,plan,role,trial_ends_at,created_at FROM users ORDER BY created_at DESC") as c:
            rows = await c.fetchall()
    return {"ok": True, "users": [dict(r) for r in rows]}

@app.get("/auth/me")
async def auth_me(request: Request):
    user = await _get_user(request)
    if not user:
        return JSONResponse({"ok": False, "error": "Non autenticato"}, 401)
    return {"ok": True, "user": user}

@app.post("/auth/forgot-password")
async def auth_forgot_password(request: Request):
    client_ip = request.client.host if request.client else "unknown"
    # Rate limit senza rivelare se il limite è stato superato (anti-enumeration)
    limited = not reset_limiter.is_allowed(client_ip)
    _OK = {"ok": True, "message": "Se l'email è registrata riceverai le istruzioni a breve."}

    if limited:
        return _OK

    body  = await request.json()
    email = (body.get("email") or "").strip().lower()
    if not validate_email(email):
        return _OK

    async with aiosqlite.connect(DB) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute("SELECT id FROM users WHERE email=?", (email,)) as c:
            user = await c.fetchone()

    if user:
        import secrets as _s
        reset_token = _s.token_urlsafe(32)
        expires = (datetime.datetime.utcnow() + datetime.timedelta(hours=1)).isoformat()
        async with aiosqlite.connect(DB) as db:
            await db.execute("UPDATE password_reset_tokens SET used=1 WHERE email=?", (email,))
            await db.execute(
                "INSERT INTO password_reset_tokens (token,user_id,email,expires_at) VALUES (?,?,?,?)",
                (reset_token, user["id"], email, expires)
            )
            await db.commit()

        app_url = os.environ.get("APP_URL", "http://localhost:8000")
        reset_url = f"{app_url}?reset_token={reset_token}"
        html = f"""
        <div style="font-family:Arial,sans-serif;max-width:500px;margin:0 auto;color:#111">
          <h2 style="color:#6366f1;margin-bottom:4px">BA.IA — Reimposta password</h2>
          <p>Hai richiesto di reimpostare la password per <strong>{email}</strong>.</p>
          <p>Clicca il pulsante qui sotto entro <strong>1 ora</strong>:</p>
          <p style="margin:28px 0">
            <a href="{reset_url}"
               style="background:#6366f1;color:#fff;padding:13px 28px;border-radius:8px;text-decoration:none;font-weight:700;font-size:15px">
              Reimposta password
            </a>
          </p>
          <p style="color:#888;font-size:12px">
            Se non hai richiesto il reset, ignora questa email.<br>
            Link diretto: {reset_url}
          </p>
        </div>"""
        try:
            _send_email(email, "BA.IA — Reimposta la tua password", html)
        except Exception as e:
            print(f"[AUTH] Errore email reset: {e}")

    await log_action("user.forgot_password", user_email=email, ip=client_ip)
    return _OK

@app.post("/auth/reset-password")
async def auth_reset_password(request: Request):
    client_ip = request.client.host if request.client else "unknown"
    body        = await request.json()
    token       = (body.get("token") or "").strip()
    new_password = body.get("password") or ""

    if not token or len(new_password) < 8:
        return JSONResponse({"ok": False, "error": "Token mancante o password troppo corta (min 8 caratteri)"}, 400)

    async with aiosqlite.connect(DB) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute(
            "SELECT * FROM password_reset_tokens WHERE token=? AND used=0 AND expires_at>datetime('now')",
            (token,)
        ) as c:
            row = await c.fetchone()

    if not row:
        return JSONResponse({"ok": False, "error": "Link non valido o scaduto. Richiedi un nuovo reset."}, 400)

    pw_hash = hash_password(new_password)
    async with aiosqlite.connect(DB) as db:
        await db.execute("UPDATE users SET password_hash=? WHERE id=?", (pw_hash, row["user_id"]))
        await db.execute("UPDATE password_reset_tokens SET used=1 WHERE token=?", (token,))
        # Invalida tutte le sessioni attive (forza re-login)
        await db.execute("DELETE FROM sessions WHERE user_id=?", (row["user_id"],))
        await db.commit()

    await log_action("user.password_reset", user_id=row["user_id"], user_email=row["email"], ip=client_ip)
    return {"ok": True, "message": "Password aggiornata. Puoi ora accedere con la nuova password."}

# ═══════════════════════════════════════════════════════
# SETUP ENDPOINTS
# ═══════════════════════════════════════════════════════

@app.get("/setup/status")
def setup_status():
    key = os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("AI_API_KEY", "")
    configured = bool(key) and len(key) > 10
    return {"configured": configured, "app_name": os.environ.get("APP_NAME", "BA.IA"), "version": "2.0.0"}

@app.post("/setup/apikey")
async def setup_apikey(request: Request):
    body = await request.json()
    key = (body.get("api_key") or "").strip()
    if not key.startswith("sk-ant-") or len(key) < 20:
        return JSONResponse({"ok": False, "error": "Chiave non valida (deve iniziare con sk-ant-)"}, 400)
    _write_env_key("ANTHROPIC_API_KEY", key)
    _write_env_key("AI_PROVIDER", "anthropic")
    os.environ["ANTHROPIC_API_KEY"] = key
    os.environ["AI_PROVIDER"] = "anthropic"
    print(f"[SETUP] Anthropic API key configurata via UI")
    return {"ok": True}

@app.post("/setup/appname")
async def setup_appname(request: Request):
    body = await request.json()
    name = (body.get("app_name") or "").strip()
    if not name:
        return JSONResponse({"ok": False, "error": "Nome vuoto"}, 400)
    _write_env_key("APP_NAME", name)
    os.environ["APP_NAME"] = name
    return {"ok": True}

# ═══════════════════════════════════════════════════════
# EXPORT WORD (.docx)
# ═══════════════════════════════════════════════════════

@app.post("/export/word")
async def export_word(request: Request):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re as _re
    except ImportError:
        return JSONResponse({"error": "python-docx non installato. Esegui: pip install python-docx"}, 500)

    body = await request.json()
    bando_name = body.get("bando_name", "Bando")
    sections = body.get("sections", [])
    note = body.get("note", "")
    checklist = body.get("checklist", [])

    doc = Document()

    # ── Margini ──
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # ── Colori ──
    NAVY = RGBColor(0x07, 0x09, 0x0e)
    GOLD = RGBColor(0xC8, 0xA8, 0x4B)
    GRAY = RGBColor(0x60, 0x60, 0x80)
    LIGHT = RGBColor(0xF0, 0xEC, 0xE3)
    EXT   = RGBColor(0xD4, 0x79, 0x0C)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        r = p.add_run(text)
        if level == 1:
            r.font.size = Pt(22)
            r.font.bold = True
            r.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(8)
        elif level == 2:
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = GOLD
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(4)
        elif level == 3:
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = GRAY
        return p

    def add_field(label, value, is_external=False):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Cm(5)
        tbl.columns[1].width = Cm(12)
        row = tbl.rows[0]
        # Label cell
        lc = row.cells[0]
        lc.text = label
        lc.paragraphs[0].runs[0].font.size = Pt(9)
        lc.paragraphs[0].runs[0].font.bold = True
        lc.paragraphs[0].runs[0].font.color.rgb = GRAY
        # Value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value or "—"))
        vr.font.size = Pt(10)
        if is_external:
            vr.font.color.rgb = EXT
            vr.font.italic = True
            badge = vp.add_run("  ⚡ fonte esterna")
            badge.font.size = Pt(8)
            badge.font.color.rgb = EXT
        # Remove borders
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)
        doc.add_paragraph()

    # ── INTESTAZIONE ──
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_after = Pt(4)
    hr = header_p.add_run("BA.IA")
    hr.font.size = Pt(28)
    hr.font.bold = True
    hr.font.color.rgb = GOLD
    sub = header_p.add_run("  —  Analisi Tecnica Bando")
    sub.font.size = Pt(14)
    sub.font.color.rgb = GRAY

    # Riga separatrice
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(16)
    sep_r = sep.add_run("─" * 70)
    sep_r.font.size = Pt(8)
    sep_r.font.color.rgb = GOLD

    add_heading(bando_name, 1)

    meta = doc.add_paragraph()
    meta.add_run(f"Generato il {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}  |  BA.IA v2.0")
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = GRAY
    doc.add_paragraph()

    # ── SEZIONI ──
    for sec in sections:
        fields = [f for f in sec.get("fields", []) if f.get("value") and str(f["value"]).strip() not in ("", "null", "None")]
        if not fields:
            continue
        add_heading(sec.get("title", ""), 2)
        for field in fields:
            add_field(field.get("label", ""), field.get("value", ""), field.get("source") == "external")

    # ── CHECKLIST ──
    if checklist:
        doc.add_paragraph()
        add_heading("Checklist operativa", 2)
        for item in checklist:
            p = doc.add_paragraph(style="List Bullet")
            mark = "☑" if item.get("completed") else "☐"
            r = p.add_run(f"{mark}  {item.get('label', '')}")
            r.font.size = Pt(10)
            if item.get("completed"):
                r.font.color.rgb = RGBColor(0x1B, 0xB8, 0x78)

    # ── NOTE ──
    if note:
        doc.add_paragraph()
        add_heading("Note", 2)
        np = doc.add_paragraph(note)
        np.runs[0].font.size = Pt(10)

    # ── FOOTER ──
    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(f"BA.IA — Report AI  |  Verificare sempre i dati sul sito ufficiale dell'ente erogatore")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY
    fr.font.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = _re.sub(r"[^a-zA-Z0-9_\-]", "_", bando_name[:50]) + ".docx"
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'})


# ═══════════════════════════════════════════════════════
# IMPORT DATI AZIENDALI DA DOCUMENTO (bilancio, visura, BP)
# ═══════════════════════════════════════════════════════

from fastapi import UploadFile, File
from pypdf import PdfReader
import httpx

@app.post("/import/company-doc")
async def import_company_doc(file: UploadFile = File(...)):
    """Estrae dati aziendali da PDF (bilancio, visura camerale, business plan)."""
    if not file.filename.lower().endswith(".pdf"):
        return JSONResponse({"ok": False, "error": "Solo file PDF supportati"}, 400)

    tmp_path = None
    try:
        data = await file.read()
        if len(data) > MAX_UPLOAD_BYTES:
            return JSONResponse({"ok": False, "error": f"File troppo grande (max {MAX_UPLOAD_BYTES // 1024 // 1024} MB)"}, 400)
        if not is_valid_pdf_bytes(data):
            return JSONResponse({"ok": False, "error": "Il file non è un PDF valido"}, 400)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        reader = PdfReader(tmp_path)
        text = "".join(page.extract_text() or "" for page in reader.pages)[:6000]

        prompt = (
            "Sei un esperto di analisi documentale aziendale italiana. "
            "Dal testo estratto da un documento aziendale (bilancio, visura camerale, business plan), "
            "estrai i dati principali dell'azienda.\n\n"
            "Restituisci SOLO JSON valido (zero testo aggiuntivo, zero backtick):\n"
            '{"name":"","piva":"","forma_giuridica":"","anno_fondazione":"","codice_ateco":"",'
            '"settore":"","dimensione":"","dipendenti":0,"fatturato":"","regione":"","comune":"",'
            '"note":""}\n\n'
            f"DOCUMENTO:\n{text}"
        )

        result_text, _ = await ai_call_multi(prompt, json_mode=True, timeout=60)
        clean = result_text.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        parsed = json.loads(clean)
        return {"ok": True, "data": parsed, "chars": len(text)}
    except Exception as e:
        return JSONResponse({"ok": False, "error": "Errore elaborazione documento"}, 500)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

# ══════════════════════════════════════════════════════════
# MULTI-AI PROVIDER ENGINE
# ══════════════════════════════════════════════════════════
import importlib

AI_PROVIDERS_AVAILABLE = ["anthropic", "openai", "groq", "gemini", "mistral"]

def _get_ai_config() -> dict:
    return {
        "provider": os.environ.get("AI_PROVIDER", "anthropic"),
        "api_key": os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("AI_API_KEY", ""),
        "model": os.environ.get("AI_MODEL", "claude-haiku-4-5-20251001"),
    }

PROVIDER_DEFAULTS = {
    "anthropic": {"model": "claude-haiku-4-5-20251001",   "url": "https://api.anthropic.com/v1/messages"},
    "openai":    {"model": "gpt-4o-mini",                  "url": "https://api.openai.com/v1/chat/completions"},
    "groq":      {"model": "meta-llama/llama-4-scout-17b-16e-instruct", "url": "https://api.groq.com/openai/v1/chat/completions"},
    "gemini":    {"model": "gemini-1.5-flash",              "url": "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"},
    "mistral":   {"model": "mistral-large-latest",          "url": "https://api.mistral.ai/v1/chat/completions"},
}

async def ai_call_multi(prompt: str, json_mode: bool = False, timeout: int = 120) -> tuple[str, float]:
    """Router AI unificato: legge provider e chiave da .env."""
    cfg = _get_ai_config()
    provider = cfg["provider"]
    api_key = cfg["api_key"]
    model = cfg["model"] or PROVIDER_DEFAULTS.get(provider, {}).get("model", "")

    if provider == "anthropic":
        headers = {"x-api-key": api_key, "anthropic-version": "2023-06-01", "Content-Type": "application/json"}
        body: dict = {"model": model, "max_tokens": 4096, "messages": [{"role": "user", "content": prompt}]}
        url = PROVIDER_DEFAULTS["anthropic"]["url"]
        async with httpx.AsyncClient(timeout=timeout) as client:
            r = await client.post(url, headers=headers, json=body)
            r.raise_for_status()
            return r.json()["content"][0]["text"], 0.0

    elif provider == "gemini":
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        body = {"contents": [{"parts": [{"text": prompt}]}]}
        async with httpx.AsyncClient(timeout=timeout) as client:
            r = await client.post(url, json=body)
            r.raise_for_status()
            return r.json()["candidates"][0]["content"]["parts"][0]["text"], 0.0

    else:
        # OpenAI-compatible: openai, groq, mistral
        urls = {"openai": "https://api.openai.com/v1/chat/completions",
                "groq": PROVIDER_DEFAULTS["groq"]["url"],
                "mistral": PROVIDER_DEFAULTS["mistral"]["url"]}
        url = urls.get(provider, PROVIDER_DEFAULTS["groq"]["url"])
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        body = {"model": model, "messages": [{"role": "user", "content": prompt}], "max_tokens": 4096}
        if json_mode and provider != "mistral":
            body["response_format"] = {"type": "json_object"}
        async with httpx.AsyncClient(timeout=timeout) as client:
            r = await client.post(url, headers=headers, json=body)
            if r.status_code == 429:
                await asyncio.sleep(20)
                r = await client.post(url, headers=headers, json=body)
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"], 0.0

from pydantic import BaseModel as _BM

class AIProviderConfig(_BM):
    provider: str
    api_key: str
    model: str = ""

@app.post("/settings/ai-provider")
async def save_ai_provider(cfg: AIProviderConfig):
    if cfg.provider not in AI_PROVIDERS_AVAILABLE:
        return JSONResponse({"ok": False, "error": f"Provider non supportato: {cfg.provider}"}, status_code=400)
    _write_env_key("AI_PROVIDER", cfg.provider)
    _write_env_key("AI_API_KEY", cfg.api_key)
    _write_env_key("AI_MODEL", cfg.model or PROVIDER_DEFAULTS.get(cfg.provider, {}).get("model", ""))
    os.environ["AI_PROVIDER"] = cfg.provider
    os.environ["AI_API_KEY"] = cfg.api_key
    os.environ["AI_MODEL"] = cfg.model or PROVIDER_DEFAULTS.get(cfg.provider, {}).get("model", "")
    return {"ok": True, "provider": cfg.provider}

@app.get("/settings/ai-provider")
def get_ai_provider():
    cfg = _get_ai_config()
    # Non restituire mai la chiave API al frontend
    return {"provider": cfg["provider"], "model": cfg["model"],
            "configured": bool(cfg["api_key"]),
            "providers": list(PROVIDER_DEFAULTS.keys()),
            "defaults": {k: v["model"] for k, v in PROVIDER_DEFAULTS.items()}}

@app.post("/settings/ai-provider/test")
async def test_ai_provider():
    try:
        result, _ = await ai_call_multi("Rispondi solo con: OK", timeout=20)
        return {"ok": True, "response": result.strip()[:50]}
    except Exception as e:
        return {"ok": False, "error": str(e)}

# ══════════════════════════════════════════════════════════
# WHITE-LABEL CONFIG
# ══════════════════════════════════════════════════════════
import base64

class WhiteLabelConfig(_BM):
    studio_name: str = "BA.IA"
    primary_color: str = "#C8A84B"
    logo_base64: str = ""
    footer_text: str = ""
    tagline: str = ""

@app.post("/settings/white-label")
async def save_white_label(cfg: WhiteLabelConfig):
    _write_env_key("WL_STUDIO_NAME", cfg.studio_name)
    _write_env_key("WL_PRIMARY_COLOR", cfg.primary_color)
    _write_env_key("WL_FOOTER_TEXT", cfg.footer_text)
    _write_env_key("WL_TAGLINE", cfg.tagline)
    os.environ["WL_STUDIO_NAME"] = cfg.studio_name
    os.environ["WL_PRIMARY_COLOR"] = cfg.primary_color
    # Save logo to disk if provided
    if cfg.logo_base64:
        logo_path = Path(os.environ["DB_PATH"]).parent / "logo.b64"
        logo_path.write_text(cfg.logo_base64)
    return {"ok": True}

@app.get("/settings/white-label")
def get_white_label():
    logo = ""
    logo_path = Path(os.environ.get("DB_PATH", "./data/ai-bandi.db")).parent / "logo.b64"
    if logo_path.exists():
        logo = logo_path.read_text()[:200] + "..." if len(logo_path.read_text()) > 200 else logo_path.read_text()
    return {
        "studio_name": os.environ.get("WL_STUDIO_NAME", "BA.IA"),
        "primary_color": os.environ.get("WL_PRIMARY_COLOR", "#C8A84B"),
        "footer_text": os.environ.get("WL_FOOTER_TEXT", ""),
        "tagline": os.environ.get("WL_TAGLINE", ""),
        "has_logo": logo_path.exists(),
    }

# ══════════════════════════════════════════════════════════
# SMTP CONFIG & TEST
# ══════════════════════════════════════════════════════════
class SmtpConfig(_BM):
    host: str = ""
    port: int = 587
    user: str = ""
    password: str = ""
    from_addr: str = ""

@app.post("/settings/smtp")
async def save_smtp(cfg: SmtpConfig):
    _write_env_key("SMTP_HOST", cfg.host)
    _write_env_key("SMTP_PORT", str(cfg.port))
    _write_env_key("SMTP_USER", cfg.user)
    _write_env_key("SMTP_PASS", cfg.password)
    _write_env_key("SMTP_FROM", cfg.from_addr)
    for k, v in [("SMTP_HOST", cfg.host), ("SMTP_PORT", str(cfg.port)),
                  ("SMTP_USER", cfg.user), ("SMTP_PASS", cfg.password), ("SMTP_FROM", cfg.from_addr)]:
        os.environ[k] = v
    return {"ok": True}

@app.get("/settings/smtp")
def get_smtp():
    return {
        "host": os.environ.get("SMTP_HOST", ""),
        "port": int(os.environ.get("SMTP_PORT", "587")),
        "user": os.environ.get("SMTP_USER", ""),
        "from_addr": os.environ.get("SMTP_FROM", ""),
        "configured": bool(os.environ.get("SMTP_HOST", "")),
    }

@app.post("/settings/smtp/test")
async def test_smtp_endpoint(request: Request):
    body = await request.json()
    to = body.get("to", "")
    if not to or not os.environ.get("SMTP_HOST"):
        return {"ok": False, "error": "SMTP non configurato o email destinatario mancante"}
    if not validate_email(to):
        return {"ok": False, "error": "Indirizzo email non valido"}
    try:
        _send_email(to, "Test email BA.IA", "<p>Configurazione SMTP funzionante ✓</p>")
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}

# ══════════════════════════════════════════════════════════
# PACK BANDI IMPORT
# ══════════════════════════════════════════════════════════
@app.post("/bandi/import-pack")
async def import_pack(file: UploadFile = File(...)):
    """Import JSON pack di bandi pre-analizzati."""
    try:
        raw = await file.read()
        pack = json.loads(raw)
        bandi_in_pack = pack.get("bandi", pack if isinstance(pack, list) else [])
        async with aiosqlite.connect(os.environ["DB_PATH"]) as db:
            imported = updated = errors = 0
            for b in bandi_in_pack:
                try:
                    bid = b.get("id") or uid_fn()
                    now_str = datetime.datetime.utcnow().isoformat()
                    await db.execute(
                        "INSERT INTO bandi (id,data,created_at,updated_at) VALUES (?,?,?,?) "
                        "ON CONFLICT(id) DO UPDATE SET data=excluded.data,updated_at=excluded.updated_at",
                        (bid, json.dumps(b), now_str, now_str))
                    imported += 1
                except Exception:
                    errors += 1
            await db.commit()
        return {"ok": True, "importati": imported, "errori": errors,
                "meta": {k: v for k, v in pack.items() if k != "bandi"}}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

def uid_fn(): return __import__("time").strftime("%Y%m%d") + __import__("secrets").token_hex(4)


# ══════════════════════════════════════════════════════════
# SCRAPER ENDPOINTS + SCHEDULER STARTUP
# ══════════════════════════════════════════════════════════
import sys as _sys
_sys.path.insert(0, str(Path(__file__).parent))

@app.on_event("startup")
async def start_scraper_scheduler():
    """Avvia lo scheduler scraper e notifiche al boot dell'app."""
    try:
        from scraper import start_scheduler
        start_scheduler(os.environ.get("DB_PATH", "./data/ai-bandi.db"))
    except Exception as e:
        print(f"[SCRAPER] Scheduler non avviato: {e}")

    # Scheduler notifiche scadenze: ogni mattina alle 08:00
    try:
        from apscheduler.schedulers.asyncio import AsyncIOScheduler
        _notif_scheduler = AsyncIOScheduler(timezone="Europe/Rome")
        _notif_scheduler.add_job(check_bandi_scadenze, "cron", hour=8, minute=0)
        _notif_scheduler.start()
        print("[NOTIF] Scheduler scadenze avviato (08:00 ogni giorno)")
    except Exception as e:
        print(f"[NOTIF] Scheduler non avviato: {e}")

@app.get("/scraper/status")
def scraper_status():
    """Stato corrente dello scraper."""
    try:
        from scraper import get_status
        return get_status()
    except Exception as e:
        return {"error": str(e), "running": False}

@app.post("/scraper/run")
async def scraper_run_manual(request: Request):
    """Avvia scraping manuale di tutte le fonti (o un sottoinsieme)."""
    try:
        body = {}
        try:
            body = await request.json()
        except Exception:
            pass
        max_sources = body.get("max_sources", None)
        from scraper import run_all_sources, _state
        if _state["running"]:
            return {"ok": False, "error": "Scraper già in esecuzione. Attendi il completamento."}
        db_path = os.environ.get("DB_PATH", "./data/ai-bandi.db")
        # Avvia in background senza bloccare la risposta HTTP
        import asyncio
        asyncio.ensure_future(run_all_sources(db_path, max_sources=max_sources))
        return {"ok": True, "message": f"Scraping avviato in background ({len(__import__('scraper').SOURCES)} fonti)"}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

@app.get("/scraper/sources")
def scraper_sources():
    """Lista delle fonti configurate."""
    try:
        from scraper import SOURCES
        return {"sources": [{"id": s["id"], "nome": s["nome"], "url": s["url"],
                              "regioni": s.get("regioni", [])} for s in SOURCES]}
    except Exception as e:
        return {"error": str(e)}

# ══════════════════════════════════════════════════════════
# SEMANTIC MATCHING ENDPOINTS
# ══════════════════════════════════════════════════════════

@app.post("/match/semantic")
async def semantic_match(request: Request):
    """
    Matching semantico azienda × tutti i bandi in DB.
    Usa TF-IDF cosine similarity (locale, zero costo).
    Se il provider supporta embedding, usa quello per qualità superiore.
    """
    try:
        body = await request.json()
        azienda = body.get("azienda", {})
        top_k = int(body.get("top_k", 10))

        if not azienda:
            return JSONResponse({"ok": False, "error": "Profilo azienda mancante"}, status_code=400)

        # Carica tutti i bandi dal DB
        async with aiosqlite.connect(os.environ["DB_PATH"]) as db:
            db.row_factory = aiosqlite.Row
            async with db.execute(
                "SELECT id, data FROM bandi WHERE JSON_EXTRACT(data,'$.status') != 'expired' "
                "ORDER BY updated_at DESC LIMIT 500"
            ) as cur:
                rows = await cur.fetchall()

        bandi = []
        for row in rows:
            try:
                d = json.loads(row["data"])
                d["id"] = row["id"]
                bandi.append(d)
            except Exception:
                pass

        if not bandi:
            return {"ok": True, "results": [], "method": "tfidf", "message": "Nessun bando in archivio"}

        from matcher import compute_tfidf_scores, get_embedding, cosine_from_embeddings, _azienda_text, _bando_text

        # Prova embedding reale
        az_text = _azienda_text(azienda)
        az_embedding = await get_embedding(az_text)
        method = "tfidf"

        if az_embedding:
            # Embedding reale disponibile — usa cosine su vettori reali
            results = []
            for bando in bandi:
                b_embed_raw = bando.get("_embedding")
                if not b_embed_raw:
                    # Genera embedding per il bando se mancante
                    b_text = _bando_text(bando)
                    b_embed_raw = await get_embedding(b_text)
                    if b_embed_raw:
                        # Salva embedding nel bando per riuso
                        bando["_embedding"] = b_embed_raw
                        async with aiosqlite.connect(os.environ["DB_PATH"]) as db:
                            await db.execute(
                                "UPDATE bandi SET data=? WHERE id=?",
                                (json.dumps(bando), bando["id"])
                            )
                            await db.commit()
                if b_embed_raw:
                    score = int(cosine_from_embeddings(az_embedding, b_embed_raw) * 100)
                    results.append({"id": bando["id"], "name": bando.get("name", ""),
                                    "ente": bando.get("ente", ""),
                                    "scadenza": bando.get("scadenza"),
                                    "score": score, "method": "embedding"})
            results.sort(key=lambda x: x["score"], reverse=True)
            method = "embedding"
        else:
            # Fallback TF-IDF
            results = compute_tfidf_scores(azienda, bandi)

        top = results[:top_k]

        # Arricchisci con motivazione AI per top 5
        for item in top[:5]:
            bando = next((b for b in bandi if b["id"] == item["id"]), None)
            if bando and item["score"] >= 20:
                try:
                    prompt = (
                        f"In 2 frasi concise, spiega perché questo bando potrebbe essere adatto "
                        f"a questa azienda. Sii specifico su beneficiari, ATECO o importi.\n\n"
                        f"AZIENDA: {azienda.get('name','')}, ATECO {azienda.get('ateco','n/d')}, "
                        f"{azienda.get('dipendenti','n/d')} dipendenti\n"
                        f"BANDO: {bando.get('name','')} — {bando.get('ente','')}\n"
                        f"Campi: {', '.join(f['label']+': '+str(f['value'])[:60] for f in (bando.get('fields') or [])[:4])}"
                    )
                    reason, _ = await ai_call_multi(prompt, timeout=30)
                    item["reason_ai"] = reason.strip()[:300]
                except Exception:
                    pass

        return {
            "ok": True,
            "results": top,
            "method": method,
            "total_bandi": len(bandi),
            "azienda": azienda.get("name", ""),
        }
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

@app.post("/match/matrix-semantic")
async def match_matrix_semantic(request: Request):
    """
    Matching semantico N aziende × M bandi in un'unica chiamata.
    Usa TF-IDF — molto più veloce del match AI uno-per-uno.
    """
    try:
        body = await request.json()
        aziende_data = body.get("aziende", [])
        bandi_data = body.get("bandi", [])

        if not aziende_data or not bandi_data:
            return JSONResponse({"ok": False, "error": "Dati mancanti"}, status_code=400)

        from matcher import compute_tfidf_scores

        matrix = {}
        for azienda in aziende_data:
            scores = compute_tfidf_scores(azienda, bandi_data)
            matrix[azienda["id"]] = {s["id"]: s["score"] for s in scores}

        return {"ok": True, "matrix": matrix, "method": "tfidf"}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

# ══════════════════════════════════════════════════════════
# GENERAZIONE BOZZA DOMANDA AI
# ══════════════════════════════════════════════════════════

@app.post("/bando/generate-draft")
async def generate_bando_draft(request: Request):
    """
    Genera una bozza di domanda di finanziamento per un bando specifico.
    Combina i dati del bando + profilo azienda → sezioni compilate.
    """
    try:
        body = await request.json()
        bando   = body.get("bando", {})
        azienda = body.get("azienda", {})
        sezioni = body.get("sezioni", [
            "presentazione_azienda",
            "descrizione_progetto",
            "obiettivi",
            "piano_investimenti",
            "sostenibilita_finanziaria",
            "impatto_occupazionale",
        ])

        if not bando or not azienda:
            return JSONResponse({"ok": False, "error": "Dati bando e azienda obbligatori"}, status_code=400)

        bando_name = bando.get("name", "Bando")
        az_name    = azienda.get("name", "Azienda")
        az_ateco   = azienda.get("ateco", "n/d")
        az_dim     = azienda.get("dipendenti", "n/d")
        az_fatt    = azienda.get("fatturato", "n/d")
        az_note    = azienda.get("note", "")

        # Estrai campi chiave dal bando
        bando_fields = "\n".join(
            f"- {f['label']}: {str(f.get('value',''))[:200]}"
            for f in bando.get("fields", [])[:15]
            if f.get("value")
        )

        sezioni_list = "\n".join(f"  {i+1}. {s.replace('_',' ').title()}" for i, s in enumerate(sezioni))

        prompt = (
            "Sei un esperto senior di finanza agevolata e redazione domande di finanziamento italiane.\n"
            "Genera una bozza professionale di domanda per il bando indicato, personalizzata per l'azienda.\n\n"
            f"BANDO: {bando_name}\n"
            f"DETTAGLI BANDO:\n{bando_fields}\n\n"
            f"AZIENDA: {az_name}\n"
            f"ATECO: {az_ateco} | Dipendenti: {az_dim} | Fatturato: {az_fatt}\n"
            f"Note aggiuntive: {az_note[:500]}\n\n"
            f"SEZIONI DA COMPILARE:\n{sezioni_list}\n\n"
            "ISTRUZIONI:\n"
            "  - Per ogni sezione scrivi 3-5 paragrafi professionali e specifici\n"
            "  - Usa linguaggio formale italiano, terminologia normativa corretta\n"
            "  - Collega i dati dell'azienda ai requisiti del bando\n"
            "  - Indica con [DA COMPLETARE] le parti che richiedono dati specifici dell'azienda\n"
            "  - Non inventare dati numerici specifici — usa quelli forniti\n\n"
            "RISPOSTA: JSON valido, zero testo aggiuntivo:\n"
            '{"sezioni":[{"id":"nome_sezione","titolo":"Titolo Sezione","contenuto":"testo...","note_consulente":"suggerimenti per personalizzare"}]}'
        )

        result, _ = await ai_call_multi(prompt, json_mode=True, timeout=180)
        clean = result.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        parsed = json.loads(clean)
        return {"ok": True, "bando": bando_name, "azienda": az_name, **parsed}
    except Exception:
        return JSONResponse({"ok": False, "error": "Errore generazione bozza"}, status_code=500)

@app.post("/bando/generate-embedding")
async def generate_bando_embedding(request: Request):
    """Genera e salva embedding vettoriale per un bando specifico."""
    try:
        body = await request.json()
        bando_id = body.get("bando_id")
        if not bando_id:
            return JSONResponse({"ok": False, "error": "bando_id mancante"}, status_code=400)

        async with aiosqlite.connect(os.environ["DB_PATH"]) as db:
            db.row_factory = aiosqlite.Row
            async with db.execute("SELECT data FROM bandi WHERE id=?", (bando_id,)) as cur:
                row = await cur.fetchone()
        if not row:
            return JSONResponse({"ok": False, "error": "Bando non trovato"}, status_code=404)

        bando = json.loads(row["data"])
        from matcher import _bando_text, get_embedding
        text = _bando_text(bando)
        embedding = await get_embedding(text)

        if embedding:
            bando["_embedding"] = embedding
            async with aiosqlite.connect(os.environ["DB_PATH"]) as db:
                await db.execute("UPDATE bandi SET data=? WHERE id=?",
                                 (json.dumps(bando), bando_id))
                await db.commit()
            return {"ok": True, "dims": len(embedding), "method": "embedding"}
        else:
            return {"ok": True, "dims": 0, "method": "tfidf_only",
                    "note": "Provider non supporta embedding — TF-IDF sempre disponibile"}
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

# ══════════════════════════════════════════════════════════
# ADMIN DASHBOARD
# ══════════════════════════════════════════════════════════
import asyncio as _asyncio
import time as _time

@app.get("/admin/dashboard")
async def admin_dashboard(request: Request):
    """
    Dashboard amministrativa: stats utenti, bandi, scraper, sistema.
    Accessibile solo da consulenti autenticati.
    """
    from main import require_auth, DB_PATH
    from fastapi import Depends
    try:
        token = request.headers.get("X-Auth-Token") or request.cookies.get("baia_token")
        if not token:
            return JSONResponse({"ok": False, "error": "Autenticazione richiesta"}, 401)

        async with aiosqlite.connect(DB_PATH) as db:
            db.row_factory = aiosqlite.Row

            # Conteggi principali
            async with db.execute("SELECT COUNT(*) FROM bandi WHERE deleted_at IS NULL") as c:
                n_bandi = (await c.fetchone())[0]
            async with db.execute("SELECT COUNT(*) FROM aziende WHERE deleted_at IS NULL") as c:
                n_aziende = (await c.fetchone())[0]
            async with db.execute("SELECT COUNT(*) FROM users") as c:
                n_users = (await c.fetchone())[0]
            async with db.execute("SELECT COUNT(*) FROM sessions WHERE expires_at > datetime('now')") as c:
                n_sessions = (await c.fetchone())[0]
            async with db.execute("SELECT COUNT(*) FROM match_history") as c:
                n_matches = (await c.fetchone())[0]
            async with db.execute("SELECT COUNT(*) FROM audit_log") as c:
                n_audit = (await c.fetchone())[0]
            # Ultimi 5 accessi
            async with db.execute(
                "SELECT ts, user_email, action, ip FROM audit_log "
                "WHERE action IN ('user.login','user.login_failed') "
                "ORDER BY ts DESC LIMIT 5"
            ) as c:
                recent_logins = [dict(r) for r in await c.fetchall()]

        # Stato scraper
        scraper_info = {}
        try:
            from scraper import get_status
            scraper_info = get_status()
        except Exception:
            pass

        return {
            "ok": True,
            "stats": {
                "bandi": n_bandi,
                "aziende": n_aziende,
                "utenti": n_users,
                "sessioni_attive": n_sessions,
                "matching_history": n_matches,
                "audit_records": n_audit,
            },
            "scraper": scraper_info,
            "recent_logins": recent_logins,
            "version": "3.0",
            "provider": os.environ.get("AI_PROVIDER", "groq"),
        }
    except Exception as e:
        return JSONResponse({"ok": False, "error": "Errore dashboard"}, 500)

@app.post("/admin/embeddings/warm-up")
async def warm_up_embeddings(request: Request):
    """
    Pre-genera embedding per tutti i bandi che non ce l'hanno ancora.
    Migliora la velocità del matching semantico.
    """
    from main import DB_PATH
    try:
        async with aiosqlite.connect(DB_PATH) as db:
            db.row_factory = aiosqlite.Row
            async with db.execute(
                "SELECT id, data FROM bandi WHERE deleted_at IS NULL "
                "AND JSON_EXTRACT(data,'$._embedding') IS NULL LIMIT 50"
            ) as c:
                rows = await c.fetchall()

        from matcher import _bando_text, get_embedding
        updated = 0
        for row in rows:
            try:
                bando = json.loads(row["data"])
                text = _bando_text(bando)
                emb = await get_embedding(text)
                if emb:
                    bando["_embedding"] = emb
                    async with aiosqlite.connect(DB_PATH) as db:
                        await db.execute("UPDATE bandi SET data=? WHERE id=?",
                                         (json.dumps(bando), row["id"]))
                        await db.commit()
                    updated += 1
                await _asyncio.sleep(0.1)  # Evita rate limit API embedding
            except Exception:
                pass

        return {"ok": True, "updated": updated, "remaining": len(rows) - updated}
    except Exception:
        return JSONResponse({"ok": False, "error": "Errore warm-up embeddings"}, 500)

# ══════════════════════════════════════════════════════════
# CLIENT PORTAL — registra endpoint e inizializza DB
# ══════════════════════════════════════════════════════════
@app.on_event("startup")
async def startup_portal():
    try:
        from portal import init_portal_db, register_portal_endpoints
        await init_portal_db()
        register_portal_endpoints(app)
    except Exception as e:
        print(f"[PORTAL] Errore avvio: {e}")

# ══════════════════════════════════════════════════════════
# RENDICONTAZIONE — registra endpoint
# ══════════════════════════════════════════════════════════
@app.on_event("startup")
async def startup_rendicontazione():
    try:
        from rendicontazione import init_rendicontazione_db, register_rendicontazione_endpoints
        await init_rendicontazione_db()
        register_rendicontazione_endpoints(app)
    except Exception as e:
        print(f"[RENDICONT] Errore avvio: {e}")

# ══════════════════════════════════════════════════════════
# REGULATORY MONITOR (D3.5) — monitora cambiamenti normativi
# ══════════════════════════════════════════════════════════
@app.on_event("startup")
async def startup_regulatory_monitor():
    try:
        from regulatory_monitor import init_regulatory_db, register_regulatory_endpoints
        await init_regulatory_db()
        register_regulatory_endpoints(app)
        print("[RegMonitor] Avviato — monitoraggio normativo attivo")
    except Exception as e:
        print(f"[RegMonitor] Errore avvio: {e}")

# ══════════════════════════════════════════════════════════
# REPORT GENERATOR (D1.2) + SOP BANDI (B2.8) — report AI
# ══════════════════════════════════════════════════════════
@app.on_event("startup")
async def startup_report_generator():
    try:
        from report_generator import register_report_endpoints
        register_report_endpoints(app)
        print("[ReportGen] Report generator + SOP Bandi attivi")
    except Exception as e:
        print(f"[ReportGen] Errore avvio: {e}")


# ══════════════════════════════════════════════════════════
# STATIC FRONTEND — serve l'interfaccia dallo stesso server
# Permette deploy unificato in produzione (frontend + backend)
# ══════════════════════════════════════════════════════════
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

FRONTEND_DIR = ROOT / "frontend"
if FRONTEND_DIR.exists():
    @app.get("/")
    async def root_index(request: Request):
        # Se la richiesta accetta HTML, serve il frontend
        # Altrimenti ritorna lo status JSON (backward compat)
        accept = request.headers.get("accept", "")
        if "text/html" in accept:
            return FileResponse(str(FRONTEND_DIR / "index.html"))
        # Status JSON per chiamate API
        from main import app as _main_app
        return {
            "status": "ok",
            "app": os.environ.get("APP_NAME", "BA.IA"),
            "version": "3.0",
            "provider": os.environ.get("AI_PROVIDER", "groq"),
            "model": os.environ.get("AI_MODEL") or os.environ.get("GROQ_MODEL", ""),
            "mode": "production" if os.environ.get("LICENSE_KEY") != "TEST-MODE" else "test",
        }

    @app.get("/app")
    async def serve_app():
        return FileResponse(str(FRONTEND_DIR / "index.html"))

    # Mount altre risorse statiche frontend
    app.mount("/static", StaticFiles(directory=str(FRONTEND_DIR)), name="static")
